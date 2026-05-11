"""
Genera BOT_ENDPOINTS_v3.docx replicando el formato visual del PDF v2.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta ────────────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x14, 0x4D, 0x5C)
NAVY_DARK = RGBColor(0x10, 0x3D, 0x49)
BLUE      = RGBColor(0x1F, 0x77, 0xB4)
TEAL      = RGBColor(0x16, 0x8A, 0x9A)
GREEN     = RGBColor(0x2E, 0xA8, 0x47)
GRAY      = RGBColor(0x55, 0x5B, 0x63)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BG_CODE   = 'F4F5F7'
BG_INFO   = 'E8F1FB'
BG_WARN   = 'FCF6E3'
BG_OK     = 'E5F4EA'
BG_ROW    = 'F8F9FA'
BG_HEAD   = '144D5C'

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

def hexstr(c): return str(c).upper()

def set_cell_bg(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        spec = kwargs.get(edge)
        if spec is None:
            continue
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), spec.get('val', 'single'))
        b.set(qn('w:sz'), str(spec.get('sz', 4)))
        b.set(qn('w:color'), spec.get('color', 'auto'))
        tcBorders.append(b)
    tcPr.append(tcBorders)

def remove_cell_borders(cell):
    set_cell_border(cell,
        top={'val':'nil'}, left={'val':'nil'},
        bottom={'val':'nil'}, right={'val':'nil'})

def add_blank(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)
    return p

def section_banner(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_bg(cell, BG_HEAD)
    remove_cell_borders(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.4)
    r = p.add_run(text)
    r.font.bold = True; r.font.size = Pt(16); r.font.color.rgb = WHITE
    add_blank(doc, 4)

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = TEAL

def body(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.bold = bold
    r.font.color.rgb = GRAY if not bold else NAVY_DARK
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.color.rgb = GRAY

def callout(doc, text, kind='info'):
    fill_map  = {'info': BG_INFO, 'warn': BG_WARN, 'ok': BG_OK}
    color_map = {'info': '1F77B4', 'warn': 'C9831B', 'ok': '2EA847'}
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = False
    cell = tbl.rows[0].cells[0]; cell.width = Inches(6.5)
    set_cell_bg(cell, fill_map[kind])
    set_cell_border(cell,
        left={'val':'single', 'sz':24, 'color': color_map[kind]},
        top={'val':'nil'}, bottom={'val':'nil'}, right={'val':'nil'})
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    r = p.add_run(text)
    r.font.size = Pt(9.5); r.font.color.rgb = NAVY_DARK
    add_blank(doc, 4)

def code_block(doc, text, lang='JSON'):
    tbl_label = doc.add_table(rows=1, cols=1); tbl_label.autofit = False
    cell = tbl_label.rows[0].cells[0]; cell.width = Inches(0.8)
    set_cell_bg(cell, BG_HEAD); remove_cell_borders(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.2)
    r = p.add_run(lang)
    r.font.bold = True; r.font.size = Pt(8); r.font.color.rgb = WHITE
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = False
    cc = tbl.rows[0].cells[0]; cc.width = Inches(6.5)
    set_cell_bg(cc, BG_CODE)
    set_cell_border(cc,
        top={'val':'single','sz':4,'color':'D5D9DE'},
        left={'val':'single','sz':4,'color':'D5D9DE'},
        bottom={'val':'single','sz':4,'color':'D5D9DE'},
        right={'val':'single','sz':4,'color':'D5D9DE'})
    first = True
    for line in text.strip('\n').split('\n'):
        p = cc.paragraphs[0] if first else cc.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.2)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Courier New'; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x1E, 0x2D, 0x40)
    add_blank(doc, 4)

def data_table(doc, headers, rows, col_widths_in=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, BG_HEAD)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            if (ri % 2) == 1: set_cell_bg(c, BG_ROW)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if ci == 0:
                r.font.color.rgb = NAVY_DARK; r.font.bold = True
    if col_widths_in:
        for row in tbl.rows:
            for ci, c in enumerate(row.cells):
                c.width = Inches(col_widths_in[ci])
    add_blank(doc, 4)

def setup_header_footer(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    hdr = sec.header
    htbl = hdr.add_table(rows=1, cols=2, width=Inches(6.5)); htbl.autofit = False
    left = htbl.rows[0].cells[0]; right = htbl.rows[0].cells[1]
    left.width = Inches(4.0); right.width = Inches(2.5)
    remove_cell_borders(left); remove_cell_borders(right)
    pl = left.paragraphs[0]
    rl = pl.add_run('DeOne Backend  •  API Endpoints Bot')
    rl.font.italic = True; rl.font.size = Pt(8.5); rl.font.color.rgb = GRAY
    pr = right.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = pr.add_run('Documento Confidencial')
    rr.font.italic = True; rr.font.size = Pt(8.5); rr.font.color.rgb = GRAY
    pPr = pl._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '8')
    top.set(qn('w:color'), '144D5C'); top.set(qn('w:space'), '4')
    pBdr.append(top); pPr.append(pBdr)

    ftr = sec.footer
    ftbl = ftr.add_table(rows=1, cols=2, width=Inches(6.5)); ftbl.autofit = False
    fl = ftbl.rows[0].cells[0]; fr = ftbl.rows[0].cells[1]
    fl.width = Inches(4.0); fr.width = Inches(2.5)
    remove_cell_borders(fl); remove_cell_borders(fr)
    pfl = fl.paragraphs[0]
    rfl = pfl.add_run('7 de mayo de 2026')
    rfl.font.italic = True; rfl.font.size = Pt(8.5); rfl.font.color.rgb = GRAY
    pfr = fr.paragraphs[0]; pfr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rfr = pfr.add_run('Página ')
    rfr.font.italic = True; rfr.font.size = Pt(8.5); rfr.font.color.rgb = GRAY
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    rEl = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '17'); rPr.append(sz)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '555B63'); rPr.append(color)
    rEl.append(rPr)
    t = OxmlElement('w:t'); t.text = '1'
    rEl.append(t)
    fld.append(rEl)
    pfr._p.append(fld)

setup_header_footer(doc)

# ── PORTADA ───────────────────────────────────────────────────────────────────
cover_top = doc.add_table(rows=1, cols=1); cover_top.autofit = False
ct = cover_top.rows[0].cells[0]; ct.width = Inches(6.5)
set_cell_bg(ct, hexstr(NAVY_DARK)); remove_cell_borders(ct)

p = ct.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('DeOne'); r.font.size = Pt(48); r.font.bold = True; r.font.color.rgb = WHITE

p = ct.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BACKEND API'); r.font.size = Pt(16); r.font.bold = True
r.font.color.rgb = RGBColor(0xB6, 0xDD, 0xE6)

p = ct.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
r = p.add_run('━━━━━━━━━━━'); r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = GREEN

p = ct.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
r = p.add_run('Documentación de Endpoints para Bot')
r.font.size = Pt(15); r.font.color.rgb = WHITE

p = ct.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Integración WhatsApp')
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xB6, 0xDD, 0xE6)

p = ct.add_paragraph(); p.paragraph_format.space_before = Pt(60)

add_blank(doc, 24)

# 3 cajas de stats
stats = [
    ('17', 'Endpoints Bot',           BLUE),
    ('2',  'Cron Jobs Automáticos',   GREEN),
    ('4',  'Módulos Principales',     NAVY_DARK),
]
stbl = doc.add_table(rows=2, cols=3); stbl.autofit = False
for i, (num, label, color) in enumerate(stats):
    top = stbl.rows[0].cells[i]; top.width = Inches(2.16)
    set_cell_bg(top, hexstr(color)); remove_cell_borders(top)
    pp = top.paragraphs[0]
    pp.paragraph_format.space_before = Pt(3); pp.paragraph_format.space_after = Pt(3)
    pp.add_run(' ')
    box = stbl.rows[1].cells[i]; box.width = Inches(2.16)
    remove_cell_borders(box)
    pn = box.paragraphs[0]; pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pn.paragraph_format.space_before = Pt(14)
    rn = pn.add_run(num)
    rn.font.size = Pt(36); rn.font.bold = True; rn.font.color.rgb = color
    pl = box.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = pl.add_run(label); rl.font.size = Pt(10); rl.font.color.rgb = GRAY
    plast = box.add_paragraph(); plast.paragraph_format.space_before = Pt(8)

add_blank(doc, 24)

info_rows = [
    ('Base URL:',      'https://prueba-supabase.onrender.com/api'),
    ('Versión:',       '3.0  —  7 de mayo de 2026'),
    ('Clasificación:', 'Documento Técnico Confidencial'),
]
itbl = doc.add_table(rows=len(info_rows), cols=2); itbl.autofit = False
for i, (k, v) in enumerate(info_rows):
    c1 = itbl.rows[i].cells[0]; c2 = itbl.rows[i].cells[1]
    c1.width = Inches(1.7); c2.width = Inches(4.8)
    remove_cell_borders(c1); remove_cell_borders(c2)
    p1 = c1.paragraphs[0]; p1.paragraph_format.left_indent = Cm(1.5)
    r1 = p1.add_run(k); r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY_DARK
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(v); r2.font.size = Pt(10); r2.font.color.rgb = GRAY

doc.add_page_break()

# ── TOC ───────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run('Tabla de Contenido')
r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = NAVY_DARK
p2 = doc.add_paragraph()
r2 = p2.add_run('━━━━━━━')
r2.font.color.rgb = GREEN; r2.font.bold = True; r2.font.size = Pt(10)
add_blank(doc, 8)

toc = [
    ('1',  'Autenticación y Formato de Respuesta', ''),
    ('2',  'Gestión de Usuarios', 'Endpoints 1-2'),
    ('3',  'Obligaciones (Compromisos Mensuales)', 'Endpoints 3-5'),
    ('4',  'Facturas', 'Endpoints 6-7'),
    ('5',  'Recargas (Depósitos de Dinero)', 'Endpoint 8'),
    ('6',  'Saldo Disponible', 'Endpoint 9'),
    ('7',  'Notificaciones  ⚠ ACTUALIZADO v3', 'Endpoints 10-12a'),
    ('8',  'Pagos', 'Endpoint 13'),
    ('9',  'Solicitudes de Recarga Automáticas', 'Endpoints 14-17'),
    ('10', 'Flujo Completo del Sistema', ''),
    ('11', 'Sistema Automático (Cron Jobs)', ''),
    ('12', 'Plantillas de Mensajes', ''),
    ('13', 'Estados del Sistema', ''),
]
ttbl = doc.add_table(rows=len(toc), cols=3); ttbl.autofit = False
for i, (n, title, sub) in enumerate(toc):
    cn = ttbl.rows[i].cells[0]; ct2 = ttbl.rows[i].cells[1]; cs = ttbl.rows[i].cells[2]
    cn.width = Inches(0.55); ct2.width = Inches(4.4); cs.width = Inches(1.55)
    remove_cell_borders(cn); remove_cell_borders(ct2); remove_cell_borders(cs)
    set_cell_bg(cn, BG_HEAD)
    pn = cn.paragraphs[0]; pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pn.paragraph_format.space_before = Pt(4); pn.paragraph_format.space_after = Pt(4)
    rn = pn.add_run(n); rn.font.bold = True; rn.font.size = Pt(11); rn.font.color.rgb = WHITE
    pt2 = ct2.paragraphs[0]
    pt2.paragraph_format.space_before = Pt(5); pt2.paragraph_format.space_after = Pt(5)
    pt2.paragraph_format.left_indent = Cm(0.4)
    rt = pt2.add_run(title); rt.font.bold = True; rt.font.size = Pt(11); rt.font.color.rgb = NAVY_DARK
    ps = cs.paragraphs[0]; ps.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ps.paragraph_format.space_before = Pt(5); ps.paragraph_format.space_after = Pt(5)
    rs = ps.add_run(sub); rs.font.italic = True; rs.font.size = Pt(9); rs.font.color.rgb = GRAY

doc.add_page_break()

# ── 1. AUTENTICACIÓN ──────────────────────────────────────────────────────────
section_banner(doc, '1. Autenticación y Formato de Respuesta')

h2(doc, 'Autenticación')
body(doc, 'Todos los endpoints requieren autenticación mediante API Key enviada en los headers HTTP.')
data_table(doc,
    ['Header', 'Valor'],
    [
        ['x-bot-api-key',                 'TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3'],
        ['x-admin-api-key (alternativo)', 'TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3'],
    ],
    [2.5, 4.0])
callout(doc, 'Ambas claves (x-bot-api-key y x-admin-api-key) usan el mismo valor y funcionan para todos los endpoints.', 'info')

body(doc, 'Si no se envía el header o el valor es incorrecto:')
code_block(doc, '''{
  "ok": false,
  "data": null,
  "error": {
    "code": "UNAUTHORIZED",
    "message": "API Key inválida o ausente"
  }
}''', 'JSON')

h2(doc, 'Formato de Respuesta Estándar')
body(doc, 'Todas las respuestas siguen exactamente este formato:')
code_block(doc, '''{
  "ok": true,
  "data": { ... },
  "error": null
}''', 'JSON')
body(doc, 'ok = true si exitosa. data contiene los datos. error es null si ok=true, o un objeto {code, message, details} si ok=false.')

h2(doc, 'Códigos de Error')
data_table(doc,
    ['HTTP', 'error.code', 'Significado'],
    [
        ['400', 'VALIDATION_ERROR',         'Datos no cumplen formato'],
        ['400', 'BAD_REQUEST',              'Error lógico (monto 0, sin facturas)'],
        ['401', 'UNAUTHORIZED',             'API Key inválida o ausente'],
        ['404', 'NOT_FOUND',                'Recurso no encontrado'],
        ['409', 'CONFLICT_DUPLICATE',       'Recurso duplicado'],
        ['409', 'INSUFFICIENT_FUNDS',       'Saldo insuficiente'],
        ['409', 'INVALID_STATE_TRANSITION', 'Estado incorrecto'],
        ['500', 'INTERNAL_ERROR',           'Error interno del servidor'],
    ],
    [0.7, 2.4, 3.4])

doc.add_page_break()

# ── 2. USUARIOS ───────────────────────────────────────────────────────────────
section_banner(doc, '2. Gestión de Usuarios')

h2(doc, 'Endpoint 1: POST /api/users/upsert')
body(doc, 'Registra un usuario nuevo o actualiza uno existente usando su teléfono como clave única. Idempotente.')
body(doc, 'Cuándo usarlo:', bold=True)
bullet(doc, 'Al inicio de cada conversación con un usuario desconocido.')
bullet(doc, 'Cuando el usuario quiere actualizar sus datos.')

body(doc, 'Parámetros del Body (JSON):', bold=True)
data_table(doc,
    ['Campo', 'Tipo', 'Req', 'Descripción'],
    [
        ['telefono', 'string', 'Sí', 'Teléfono del usuario (mín 7 chars). Clave única'],
        ['nombre',   'string', 'No', 'Nombre del usuario'],
        ['apellido', 'string', 'No', 'Apellido del usuario'],
        ['correo',   'string', 'No', 'Email válido del usuario'],
    ],
    [1.4, 0.9, 0.6, 3.6])

body(doc, 'Ejemplo de Request:', bold=True)
code_block(doc, '''POST /api/users/upsert
Content-Type: application/json
x-bot-api-key: TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3

{
  "telefono": "573046757626",
  "nombre": "Laura",
  "apellido": "Duran",
  "correo": "laura@email.com"
}''', 'HTTP')

body(doc, 'Respuesta Exitosa (201 si nuevo, 200 si existente):', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": {
    "usuario_id": "49f3c602-80c8-4c59-9ee6-a005bbb86f08",
    "creado": true
  },
  "error": null
}''', 'JSON')
callout(doc, 'creado: true = usuario nuevo, false = actualizado. Guardar usuario_id para referencias internas.', 'info')

h2(doc, 'Endpoint 2: PUT /api/users/plan')
body(doc, 'Asigna o cambia el plan de pago del usuario. Debe llamarse ANTES de generar solicitudes de recarga (endpoint 14).')
body(doc, 'Planes Disponibles:', bold=True)
data_table(doc,
    ['Plan', 'Recargas/Mes', 'Distribución'],
    [
        ['control',      '1', 'Día 1: pago total'],
        ['tranquilidad', '2', 'Día 1 y 15: divide por vencimiento'],
        ['respaldo',     '2', 'Día 1 y 15: divide por vencimiento'],
    ],
    [1.6, 1.4, 3.5])

body(doc, 'Request:', bold=True)
code_block(doc, '''PUT /api/users/plan
x-bot-api-key: TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3

{ "telefono": "573046757626", "plan": "tranquilidad" }''', 'HTTP')

body(doc, 'Respuesta Exitosa (200):', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": {
    "usuario_id": "49f3c602-80c8-4c59-9ee6-a005bbb86f08",
    "telefono": "573046757626",
    "plan_anterior": "control",
    "plan_nuevo": "tranquilidad"
  },
  "error": null
}''', 'JSON')

doc.add_page_break()

# ── 3. OBLIGACIONES ───────────────────────────────────────────────────────────
section_banner(doc, '3. Obligaciones (Compromisos Mensuales)')

h2(doc, 'Endpoint 3: POST /api/obligaciones')
body(doc, 'Crea una obligación mensual: contenedor que agrupa todas las facturas que el usuario debe pagar en un período (mes).')
body(doc, 'Parámetros del Body:', bold=True)
data_table(doc,
    ['Campo', 'Tipo', 'Req', 'Descripción'],
    [
        ['telefono',    'string', 'Sí', 'Teléfono del usuario'],
        ['descripcion', 'string', 'Sí', 'Ej: "Pagos de Mayo 2026"'],
        ['periodo',     'string', 'Sí', 'Fecha YYYY-MM-DD (se normaliza al día 1)'],
    ],
    [1.4, 0.9, 0.6, 3.6])

body(doc, 'Request:', bold=True)
code_block(doc, '''POST /api/obligaciones
x-bot-api-key: TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3

{
  "telefono": "573046757626",
  "descripcion": "Pagos de Mayo 2026",
  "periodo": "2026-05-01"
}''', 'HTTP')

body(doc, 'Respuesta Exitosa (201):', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": {
    "id": "86a0c709-3ca9-41bc-9106-226cac7cf4ba",
    "usuario_id": "49f3c602-80c8-4c59-9ee6-a005bbb86f08",
    "descripcion": "Pagos de Mayo 2026",
    "periodo": "2026-05-01",
    "estado": "activa",
    "total_facturas": 0,
    "facturas_pagadas": 0,
    "monto_total": 0,
    "monto_pagado": 0,
    "creado_en": "2026-05-07T10:30:00.000Z"
  },
  "error": null
}''', 'JSON')
callout(doc, 'Guardar el campo "id" de la respuesta. Se necesita para asociar facturas (endpoint 6) y generar solicitudes de recarga (endpoint 14).', 'warn')

h2(doc, 'Endpoint 4: GET /api/obligaciones?telefono=XXX')
body(doc, 'Devuelve TODAS las obligaciones del usuario con su array de facturas, montos y porcentaje de progreso.')
data_table(doc,
    ['Param', 'Tipo', 'Req', 'Descripción'],
    [
        ['telefono', 'string', 'Sí', 'Teléfono del usuario'],
        ['estado',   'string', 'No', '"activa", "en_progreso", "completada", "cancelada"'],
    ],
    [1.4, 0.9, 0.6, 3.6])

body(doc, 'Respuesta Exitosa (200):', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": [
    {
      "id": "86a0c709-...",
      "descripcion": "Pagos de Mayo 2026",
      "periodo": "2026-05-01",
      "estado": "activa",
      "total_facturas": 3,
      "facturas_pagadas": 1,
      "monto_total": 250000,
      "monto_pagado": 85000,
      "progreso": 33,
      "facturas": [
        {
          "id": "18e6dcfd-...",
          "servicio": "Internet ETB",
          "monto": 85000,
          "estado": "pagada",
          "validacion_estado": "validada",
          "etiqueta": "internet",
          "fecha_vencimiento": "2026-05-10"
        }
      ]
    }
  ],
  "error": null
}''', 'JSON')

h2(doc, 'Endpoint 5: GET /api/obligaciones/:id')
body(doc, 'Detalle completo de UNA obligación específica con todas sus facturas y datos del usuario (nombre, teléfono).')
body(doc, 'Misma estructura que endpoint 4 pero un solo objeto en vez de array. Incluye el campo "usuarios" con nombre, apellido y teléfono del dueño.')

doc.add_page_break()

# ── 4. FACTURAS ───────────────────────────────────────────────────────────────
section_banner(doc, '4. Facturas')

h2(doc, 'Endpoint 6: POST /api/facturas/captura')
body(doc, 'Registra una factura dentro de una obligación. El bot extrae los datos (foto, PDF, texto) y los envía aquí.')
callout(doc, 'CAMBIO v3 — La factura ahora se crea con estado="pendiente" y validacion_estado="sin_validar". El admin la cambia a validacion_estado="validada" antes de poder pagarse.', 'warn')

body(doc, 'Parámetros del Body:', bold=True)
data_table(doc,
    ['Campo', 'Tipo', 'Req', 'Nullable', 'Descripción'],
    [
        ['telefono',             'string', 'Sí', 'No', 'Teléfono del usuario'],
        ['obligacion_id',        'UUID',   'Sí', 'No', 'ID de la obligación'],
        ['servicio',             'string', 'Sí', 'No', 'Nombre del servicio'],
        ['monto',                'number', 'No', 'Sí', 'Valor en pesos. null = revisión'],
        ['fecha_vencimiento',    'string', 'No', 'Sí', 'Fecha límite YYYY-MM-DD'],
        ['fecha_emision',        'string', 'No', 'Sí', 'Fecha emisión YYYY-MM-DD'],
        ['referencia_pago',      'string', 'No', 'Sí', 'Referencia de factura'],
        ['etiqueta',             'string', 'No', 'Sí', '"internet", "gas", etc.'],
        ['origen',               'string', 'No', 'Sí', '"bot_whatsapp", "manual", "ocr"'],
        ['archivo_url',          'string', 'No', 'Sí', 'URL imagen/PDF'],
        ['extraccion_estado',    'string', 'No', 'No', '"ok", "dudosa", "fallida"'],
        ['extraccion_confianza', 'number', 'No', 'Sí', 'Confianza 0.0 a 1.0'],
    ],
    [1.6, 0.7, 0.5, 0.7, 3.0])

body(doc, 'Ejemplo — Factura completa:', bold=True)
code_block(doc, '''{
  "telefono": "573046757626",
  "obligacion_id": "86a0c709-3ca9-41bc-9106-226cac7cf4ba",
  "servicio": "Vanti S.A. ESP",
  "monto": 50950,
  "fecha_vencimiento": "2026-05-24",
  "fecha_emision": "2026-05-01",
  "referencia_pago": "7890123456",
  "etiqueta": "gas",
  "origen": "bot_whatsapp",
  "archivo_url": "https://storage.supabase.co/facturas/gas_may.pdf",
  "extraccion_estado": "ok",
  "extraccion_confianza": 0.95
}''', 'JSON')

body(doc, 'Respuesta Exitosa (201):', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": {
    "factura_id": "63dd4b3b-9e6c-43b1-a6b6-ff5858554733",
    "estado": "pendiente",
    "validacion_estado": "sin_validar",
    "requiere_revision": false
  },
  "error": null
}''', 'JSON')

body(doc, 'Estados posibles de una factura (v3):', bold=True)
data_table(doc,
    ['estado', 'validacion_estado', 'Significado'],
    [
        ['pendiente', 'sin_validar', 'Recién capturada. Esperando revisión admin'],
        ['pendiente', 'validada',    'Admin confirmó. Lista para generar pago'],
        ['pendiente', 'rechazada',   'Admin rechazó esta factura'],
        ['pagada',    'validada',    'Pago ejecutado exitosamente'],
        ['anulada',   '—',           'Factura cancelada'],
    ],
    [1.3, 1.7, 3.4])

h2(doc, 'Endpoint 7: GET /api/facturas/obligacion/:obligacionId')
body(doc, 'Lista todas las facturas de una obligación con todos los campos.')
body(doc, 'Respuesta Exitosa (200):', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": [
    {
      "referencia_pago": "ETB-2026-001",
      "servicio": "Internet ETB",
      "monto": 85000,
      "estado": "pendiente",
      "validacion_estado": "validada",
      "origen": "bot_whatsapp",
      "etiqueta": "internet",
      "fecha_emision": "2026-04-25",
      "fecha_vencimiento": "2026-05-10",
      "periodo": "2026-05-01",
      "_id": "63dd4b3b-9e6c-43b1-a6b6-ff5858554733"
    }
  ],
  "error": null
}''', 'JSON')
callout(doc, 'El campo _id (UUID interno) es el que se debe usar en el endpoint de pagos (13) como factura_id.', 'info')

doc.add_page_break()

# ── 5. RECARGAS ───────────────────────────────────────────────────────────────
section_banner(doc, '5. Recargas (Depósitos de Dinero)')

h2(doc, 'Endpoint 8: POST /api/recargas/reportar')
body(doc, 'Registra que el usuario depositó dinero (Nequi, PSE, etc.). Queda en "en_validacion" hasta que un admin la apruebe.')
callout(doc, 'Si se envía referencia_tx y ya existe una recarga con esa referencia, NO se duplica — devuelve la existente. La recarga NO está disponible inmediatamente; debe ser aprobada por admin.', 'warn')

data_table(doc,
    ['Campo', 'Tipo', 'Req', 'Descripción'],
    [
        ['telefono',        'string', 'Sí', 'Teléfono del usuario'],
        ['periodo',         'string', 'Sí', 'Periodo YYYY-MM-DD'],
        ['monto',           'number', 'Sí', 'Monto recargado (positivo)'],
        ['comprobante_url', 'string', 'Sí', 'URL imagen del comprobante'],
        ['referencia_tx',   'string', 'No', 'Referencia bancaria. Previene duplicados'],
    ],
    [1.6, 0.9, 0.6, 3.4])

body(doc, 'Request:', bold=True)
code_block(doc, '''POST /api/recargas/reportar
x-bot-api-key: TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3

{
  "telefono": "573046757626",
  "periodo": "2026-05-01",
  "monto": 130000,
  "comprobante_url": "https://storage.supabase.co/comprobantes/nequi_001.jpg",
  "referencia_tx": "NEQUI-2026050712345"
}''', 'HTTP')

body(doc, 'Respuesta Exitosa (201):', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": {
    "recarga_id": "974bad6d-a234-4b56-c789-d012ef345678",
    "estado": "en_validacion"
  },
  "error": null
}''', 'JSON')

body(doc, 'Si ya existía esa referencia_tx (200 — idempotente):', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": {
    "recarga_id": "974bad6d-a234-4b56-c789-d012ef345678",
    "estado": "en_validacion",
    "mensaje": "Recarga ya reportada con esta referencia de transacción"
  },
  "error": null
}''', 'JSON')

doc.add_page_break()

# ── 6. SALDO ──────────────────────────────────────────────────────────────────
section_banner(doc, '6. Saldo Disponible')

h2(doc, 'Endpoint 9: GET /api/disponible')
body(doc, 'Calcula en tiempo real cuánto dinero tiene disponible el usuario.')
body(doc, 'Fórmula: disponible = total_recargas_aprobadas − total_pagos_realizados', bold=True)

data_table(doc,
    ['Param', 'Tipo', 'Req', 'Descripción'],
    [
        ['telefono', 'string', 'Sí', 'Teléfono del usuario'],
        ['periodo',  'string', 'Sí', 'Periodo YYYY-MM-DD'],
    ],
    [1.4, 0.9, 0.6, 3.6])

body(doc, 'Ejemplo:', bold=True)
code_block(doc, 'GET /api/disponible?telefono=573046757626&periodo=2026-05-01', 'HTTP')

body(doc, 'Respuesta Exitosa (200):', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": {
    "usuario_id": "49f3c602-80c8-4c59-9ee6-a005bbb86f08",
    "periodo": "2026-05-01",
    "total_recargas_aprobadas": 250000,
    "total_pagos_pagados": 130000,
    "disponible": 120000
  },
  "error": null
}''', 'JSON')

doc.add_page_break()

# ── 7. NOTIFICACIONES ─────────────────────────────────────────────────────────
section_banner(doc, '7. Notificaciones  ⚠ ACTUALIZADO v3')

callout(doc,
    'CAMBIO CRÍTICO v3 — Todos los tipos antiguos (solicitud_recarga_inicio_mes, recordatorio_recarga, recarga_confirmada, '
    'recarga_aprobada, pago_confirmado, factura_validada, obligacion_completada, nueva_obligacion) fueron ELIMINADOS. '
    'El bot ahora SOLO recibe 3 tipos.', 'warn')

body(doc, 'Tipos válidos que recibe el bot:', bold=True)
data_table(doc,
    ['Tipo', 'Cuándo se genera'],
    [
        ['solicitud_recarga',           'Cron cada 30 min detecta obligaciones activas sin saldo suficiente'],
        ['obligacion_cumplida',         'Pago de UNA factura (sin otros pagos en ventana de 30 min)'],
        ['obligaciones_pagadas_grupal', '2+ pagos dentro de ventana de 30 min (tipo virtual creado en consumo)'],
    ],
    [2.4, 4.1])

h2(doc, 'Endpoint 10: GET /api/notificaciones/pendientes/:telefono')
body(doc, 'Devuelve notificaciones pendientes para enviar al usuario por WhatsApp. Comportamiento atómico: al consultar, '
          'automáticamente marca todas como "enviada". No es necesario llamar 11 ni 12 después.')
callout(doc, '[NUEVO v3] La agrupación de pagos sucede aquí: si hay ≥2 pagos con creado_en a ≤30 minutos, se devuelven como '
             'una sola notificación tipo "obligaciones_pagadas_grupal" con payload.obligaciones[].', 'ok')

body(doc, 'Ejemplo:', bold=True)
code_block(doc, 'GET /api/notificaciones/pendientes/573046757626\nx-bot-api-key: TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3', 'HTTP')

body(doc, 'Respuesta Exitosa (200) — múltiples tipos posibles:', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": [
    {
      "tipo": "solicitud_recarga",
      "payload": {
        "valor_recarga": 215000,
        "mensaje": "Laura 👋\\n\\nEs momento de recargar tu cuenta..."
      }
    },
    {
      "tipo": "obligacion_cumplida",
      "payload": { "etiqueta": "Internet ETB", "monto_total": 85000 }
    },
    {
      "tipo": "obligaciones_pagadas_grupal",
      "payload": {
        "obligaciones": [
          { "etiqueta": "Gas Vanti",   "valor": 50950 },
          { "etiqueta": "Energía EPM", "valor": 120000 }
        ]
      }
    }
  ],
  "error": null
}''', 'JSON')

body(doc, 'Flujo recomendado para el bot:', bold=True)
code_block(doc, '''1. GET /api/notificaciones/pendientes/{telefono}
   -> Las notificaciones se marcan como "enviada" AUTOMÁTICAMENTE.

2. Por cada notificación en data[]:
   a. tipo = "solicitud_recarga"
      -> Leer payload.mensaje (texto listo) o construir desde payload.valor_recarga.

   b. tipo = "obligacion_cumplida"
      -> Usar payload.etiqueta y payload.monto_total.

   c. tipo = "obligaciones_pagadas_grupal"
      -> Iterar payload.obligaciones[] -> { etiqueta, valor }
      -> Construir mensaje con la lista de facturas pagadas.

3. NO llamar EP11 ni EP12 después — ya están marcadas.''', 'TEXT')

h2(doc, 'Endpoint 11: PUT /api/notificaciones/:id  (opcional)')
body(doc, 'Actualiza el estado de UNA notificación. Útil para reportar fallos al enviar por WhatsApp.')
code_block(doc, '''PUT /api/notificaciones/7cc2d2eb-6b6b-4d6e-b3c2-2e9e4e6ed7ca
x-bot-api-key: TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3

{ "estado": "fallida", "ultimo_error": "WhatsApp API timeout después de 30s" }''', 'HTTP')

h2(doc, 'Endpoint 12: POST /api/notificaciones/batch-enviadas  (opcional)')
body(doc, 'Marca múltiples notificaciones como "enviada" en una sola llamada.')
code_block(doc, '''POST /api/notificaciones/batch-enviadas
x-bot-api-key: TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3

{ "ids": ["7cc2d2eb-...", "a1b2c3d4-..."] }''', 'HTTP')

h2(doc, 'Endpoint 12a: GET /api/notificaciones/pendientes-hoy')
body(doc, 'Endpoint estratégico para job global. Devuelve TODAS las notificaciones tipo "solicitud_recarga" creadas HOY '
          'en estado pendiente para TODOS los usuarios. Marca automáticamente como "entregada" en la misma consulta.')
callout(doc, 'CAMBIO v3 — Antes filtraba "solicitud_recarga_inicio_mes". Ahora filtra "solicitud_recarga". Idempotente: '
             '2da consulta del mismo día devuelve total: 0.', 'warn')

body(doc, 'Respuesta (200) — incluye datos del usuario en cada item:', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": {
    "total": 2,
    "notificaciones": [
      {
        "id": "e4c9a1f2-...",
        "tipo": "solicitud_recarga",
        "estado": "pendiente",
        "usuarios": {
          "nombre": "Laura",
          "apellido": "Duran",
          "telefono": "573046757626"
        },
        "payload": {
          "valor_recarga": 215000,
          "mensaje": "Hola Laura...\\nEs momento de recargar..."
        }
      }
    ]
  },
  "error": null
}''', 'JSON')

body(doc, 'Garantías de comportamiento:', bold=True)
bullet(doc, 'Filtra SOLO "solicitud_recarga" (no más _inicio_mes).')
bullet(doc, 'Solo notificaciones del día actual en estado "pendiente".')
bullet(doc, 'Marca TODAS como "enviada" automáticamente.')
bullet(doc, 'Idempotente: segunda consulta del mismo día devuelve [].')
bullet(doc, 'Si el job falla, las notificaciones permanecen en "pendiente".')

doc.add_page_break()

# ── 8. PAGOS ──────────────────────────────────────────────────────────────────
section_banner(doc, '8. Pagos')

h2(doc, 'Endpoint 13: POST /api/pagos/crear')
body(doc, 'Crea un pago para una factura específica. Verifica que:')
bullet(doc, 'La factura exista y tenga validacion_estado = "validada".')
bullet(doc, 'El usuario tenga saldo disponible suficiente.')
body(doc, 'Si ambas se cumplen, crea el pago en estado "en_proceso". El monto se descuenta del saldo.')

data_table(doc,
    ['Campo', 'Tipo', 'Req', 'Descripción'],
    [
        ['telefono',   'string', 'Sí', 'Teléfono del usuario'],
        ['factura_id', 'UUID',   'Sí', 'UUID de la factura (campo _id del endpoint 7)'],
    ],
    [1.4, 0.9, 0.6, 3.6])

body(doc, 'Request:', bold=True)
code_block(doc, '''POST /api/pagos/crear
x-bot-api-key: TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3

{
  "telefono": "573046757626",
  "factura_id": "63dd4b3b-9e6c-43b1-a6b6-ff5858554733"
}''', 'HTTP')

body(doc, 'Respuesta Exitosa (201):', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": {
    "pago_id": "0da485dd-f234-5678-9abc-def012345678",
    "estado": "en_proceso",
    "monto": 85000,
    "servicio": "Internet ETB"
  },
  "error": null
}''', 'JSON')

body(doc, 'Errores posibles:', bold=True)
data_table(doc,
    ['Código', 'Mensaje'],
    [
        ['NOT_FOUND',                'Factura no encontrada'],
        ['INVALID_STATE_TRANSITION', "Factura con validacion_estado='sin_validar'. Debe estar 'validada'."],
        ['INSUFFICIENT_FUNDS',       'Fondos insuficientes. Disponible: $X, Requerido: $Y'],
    ],
    [2.2, 4.3])

doc.add_page_break()

# ── 9. SOLICITUDES ────────────────────────────────────────────────────────────
section_banner(doc, '9. Solicitudes de Recarga Automáticas')

h2(doc, 'Endpoint 14: POST /api/solicitudes-recarga/generar')
body(doc, 'Analiza las facturas de una obligación y genera solicitudes de recarga según el plan del usuario. '
          'Requiere facturas con validacion_estado = "validada".')

callout(doc, 'Plan CONTROL (1 cuota): 1 solicitud por monto total. Fecha límite = vencimiento más próximo. Recordatorio = 5 días antes.', 'info')
callout(doc, 'Plan TRANQUILIDAD / RESPALDO (2 cuotas): divide facturas en 2 grupos por fecha de vencimiento. '
             'Cuota 1: vencen días 1-15. Cuota 2: vencen días 16-31. Si todas caen en la misma mitad, divide 50/50.', 'info')

body(doc, 'Request:', bold=True)
code_block(doc, '''POST /api/solicitudes-recarga/generar
x-bot-api-key: TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3

{
  "telefono": "573046757626",
  "obligacion_id": "260b7859-a392-45bc-95c9-5a7f627a93f7"
}''', 'HTTP')

body(doc, 'Respuesta Exitosa (201) — Plan tranquilidad:', bold=True)
code_block(doc, '''{
  "ok": true,
  "data": {
    "solicitudes": [
      {
        "id": "09e2ed08-...",
        "numero_cuota": 1,
        "total_cuotas": 2,
        "monto_solicitado": 215000,
        "fecha_limite": "2026-05-10",
        "fecha_recordatorio": "2026-05-05",
        "estado": "pendiente"
      },
      {
        "id": "3e30936d-...",
        "numero_cuota": 2,
        "total_cuotas": 2,
        "monto_solicitado": 120000,
        "fecha_limite": "2026-05-22",
        "fecha_recordatorio": "2026-05-17",
        "estado": "pendiente"
      }
    ],
    "plan": "tranquilidad",
    "monto_total": 335000,
    "total_cuotas": 2
  },
  "error": null
}''', 'JSON')

h2(doc, 'Endpoint 15: GET /api/solicitudes-recarga')
body(doc, 'Lista todas las solicitudes de recarga de un usuario. Permite filtrar por estado y obligación.')
data_table(doc,
    ['Param', 'Tipo', 'Req', 'Descripción'],
    [
        ['telefono',      'string', 'Sí', 'Teléfono del usuario'],
        ['estado',        'string', 'No', '"pendiente", "parcial", "cumplida", "vencida", "cancelada"'],
        ['obligacion_id', 'UUID',   'No', 'Filtrar por obligación'],
    ],
    [1.6, 0.9, 0.6, 3.4])

h2(doc, 'Endpoint 16: POST /api/solicitudes-recarga/verificar-recordatorios')
body(doc, 'Busca solicitudes próximas a vencer, verifica si el usuario tiene saldo, y genera notificación '
          '"solicitud_recarga" si no tiene fondos.')
callout(doc, 'El cron job cada 30 minutos ya ejecuta esta verificación automáticamente para TODOS los usuarios. '
             'Este endpoint sirve como respaldo o para verificar un usuario específico.', 'info')

code_block(doc, '''POST /api/solicitudes-recarga/verificar-recordatorios
{ "telefono": "573046757626" }

// Respuesta (200):
{ "ok": true, "data": { "recordatorios_generados": 1, "detalle": [...] } }''', 'JSON')

h2(doc, 'Endpoint 17: PUT /api/solicitudes-recarga/:id/fechas')
body(doc, 'Permite al usuario cambiar las fechas límite de sus cuotas. Recalcula automáticamente fecha_recordatorio '
          '(5 días antes) y resetea recordatorio_enviado.')
callout(doc, 'Debe enviarse al menos una fecha. Solo se pueden modificar solicitudes en estado "pendiente" o "parcial".', 'warn')
code_block(doc, '''PUT /api/solicitudes-recarga/09e2ed08-.../fechas
x-bot-api-key: TK2026A7F9X3M8N2P5Q1R4T6Y8U0I9O3

{ "fecha_cuota_1": "2026-05-05", "fecha_cuota_2": "2026-05-18" }''', 'HTTP')

doc.add_page_break()

# ── 10. FLUJO ─────────────────────────────────────────────────────────────────
section_banner(doc, '10. Flujo Completo del Sistema')

def fase_banner(doc, label, right):
    tbl = doc.add_table(rows=1, cols=2); tbl.autofit = False
    cl = tbl.rows[0].cells[0]; cr = tbl.rows[0].cells[1]
    cl.width = Inches(4.0); cr.width = Inches(2.5)
    set_cell_bg(cl, BG_HEAD); set_cell_bg(cr, BG_HEAD)
    remove_cell_borders(cl); remove_cell_borders(cr)
    pl = cl.paragraphs[0]
    pl.paragraph_format.space_before = Pt(4); pl.paragraph_format.space_after = Pt(4)
    pl.paragraph_format.left_indent = Cm(0.4)
    rl = pl.add_run(label)
    rl.font.bold = True; rl.font.size = Pt(11); rl.font.color.rgb = WHITE
    pr = cr.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.space_before = Pt(4); pr.paragraph_format.space_after = Pt(4)
    rr = pr.add_run(right)
    rr.font.italic = True; rr.font.size = Pt(10); rr.font.color.rgb = RGBColor(0xB6, 0xDD, 0xE6)
    add_blank(doc, 2)

fase_banner(doc, 'FASE 1: REGISTRO Y CONFIGURACIÓN', 'Bot ↔ Usuario')
for line in [
    '1. Usuario llega por WhatsApp     → POST /api/users/upsert',
    '2. Elige plan                     → PUT /api/users/plan',
    '3. Crear obligación del mes       → POST /api/obligaciones',
    '4. Cargar facturas (foto/PDF)     → POST /api/facturas/captura  (por cada factura)',
    '5. Admin valida facturas          → validacion_estado = "validada"',
    '6. Generar solicitudes            → POST /api/solicitudes-recarga/generar',
]: bullet(doc, line)

fase_banner(doc, 'FASE 2: EVALUACIÓN AUTOMÁTICA', 'Cron Jobs')
for line in [
    'Cron cada 30 minutos (*/30 * * * *) — jobEvaluacionRecargas:',
    '   • Evalúa obligaciones activas, calcula montos, verifica saldo.',
    '   • Genera notificación "solicitud_recarga" si no hay fondos suficientes.',
    '   • No genera duplicados si ya existe una del mismo día.',
    'Cron cada 6 horas (0 */6 * * *) — jobVerificarInactividad:',
    '   • Detecta usuarios sin respuesta 24-48h.',
    '   • Crea alerta interna para admin (el bot no la recibe).',
]: bullet(doc, line)

fase_banner(doc, 'FASE 3: RECARGA Y APROBACIÓN', 'Bot ↔ Usuario ↔ Admin')
for line in [
    '7. Bot entrega notificaciones     → GET /api/notificaciones/pendientes/:tel',
    '8. Usuario reporta recarga        → POST /api/recargas/reportar',
    '9. Admin aprueba recarga          → cancela solicitudes cubiertas, actualiza saldo',
]: bullet(doc, line)

fase_banner(doc, 'FASE 4: PAGOS Y SEGUIMIENTO', 'Bot + Admin')
for line in [
    '10. Verificar saldo               → GET /api/disponible',
    '11. Crear pago por factura        → POST /api/pagos/crear  (requiere validacion_estado=validada)',
    '12. Consultar estado              → GET /api/obligaciones, /facturas, /solicitudes-recarga',
    '13. Bot recibe notif de pago      → obligacion_cumplida (1) o obligaciones_pagadas_grupal (≥2 en <30 min)',
]: bullet(doc, line)

doc.add_page_break()

# ── 11. CRON JOBS ─────────────────────────────────────────────────────────────
section_banner(doc, '11. Sistema Automático de Evaluación (Cron Jobs)')
body(doc, 'El servidor ejecuta cron jobs automáticamente que evalúan obligaciones, recalculan solicitudes y generan notificaciones.')

h2(doc, 'Job 1: jobEvaluacionRecargas — Cada 30 minutos')
body(doc, 'Archivo: src/jobs/recordatorios.job.js')
body(doc, 'Programación: cron.schedule("*/30 * * * *", ...)', bold=True)
code_block(doc, '''Secuencia de ejecución:
1. obtenerObligacionesActivas()    -> Lista de obligaciones activas
2. FOR cada obligación:
     a. evaluarObligacion()         -> Calcula monto pendiente
     b. existeNotificacionHoy()     -> Evita duplicados diarios
     c. prepararDatosNotificacion() -> Construye payload
     d. crearNotificacionRecarga()  -> Inserta notificación tipo "solicitud_recarga"
3. Retorna resumen: procesadas, creadas, errores''', 'TEXT')

callout(doc, 'CAMBIO v3 — La frecuencia ahora es cada 30 minutos (antes era 9:00 AM diario). El tipo único de '
             'notificación generada es "solicitud_recarga".', 'warn')

h2(doc, 'Job 2: jobVerificarInactividad — Cada 6 horas')
body(doc, 'Archivo: src/modules/notificaciones/notificaciones.service.js')
body(doc, 'Programación: cron.schedule("0 */6 * * *", ...)', bold=True)
code_block(doc, '''Lógica de detección:
1. Rango: notificaciones enviadas 24-48h atrás
2. Tipo: solicitud_recarga
3. Estado: enviada
4. Verifica: ¿Ya existe alerta para esta notificación? -> omitir si sí
5. Verifica: ¿Hay recargas después del mensaje?
     -> Si NO -> crearAlertaAdmin() (no llega al bot, solo al admin)''', 'TEXT')

body(doc, 'Protecciones del sistema:', bold=True)
bullet(doc, 'No genera notificaciones duplicadas el mismo día.')
bullet(doc, 'Control de concurrencia (un solo job a la vez).')
bullet(doc, 'Si node-cron no está instalado, el servidor arranca sin jobs.')
bullet(doc, 'Los montos se recalculan automáticamente al cambiar facturas.')

doc.add_page_break()

# ── 12. PLANTILLAS ────────────────────────────────────────────────────────────
section_banner(doc, '12. Plantillas de Mensajes Automáticos')
body(doc, 'Para "solicitud_recarga", payload.mensaje viene con texto listo. Para los otros 2 tipos, el bot construye el mensaje.')

h2(doc, 'Tipo: solicitud_recarga')
code_block(doc, '''{nombre} 👋

Es momento de recargar tu cuenta para cubrir tus próximas obligaciones 🙌

Tu saldo actual en deOne es de $X
Valor a recargar: $Y

Puedes hacer la recarga a la llave 0090944088.
Cuando la hagas, envíame el comprobante y yo me encargo del resto deOne 👍''', 'WhatsApp')

h2(doc, 'Tipo: obligacion_cumplida  (1 sola factura)')
code_block(doc, '''¡{nombre}! 🙌
Ya hice el pago de {etiqueta} por ${valor}.

Tu saldo actualizado en deOne es de $X

El comprobante ya quedó cargado en tu enlace habitual.''', 'WhatsApp')

h2(doc, 'Tipo: obligaciones_pagadas_grupal  (varias facturas)')
body(doc, 'Iterar payload.obligaciones[] → { etiqueta, valor }:')
code_block(doc, '''¡{nombre}! 🙌

Ya hice el pago de:
{etiqueta_1} por ${valor_1}.
{etiqueta_2} por ${valor_2}.
...

Tu saldo actualizado en deOne es de $X

¡Los comprobantes ya quedaron cargados en tu enlace habitual!''', 'WhatsApp')

callout(doc, 'Llave de recarga: 0090944088 — Esta llave aparece en los mensajes automáticos.', 'warn')

doc.add_page_break()

# ── 13. ESTADOS ───────────────────────────────────────────────────────────────
section_banner(doc, '13. Resumen de Estados del Sistema')

for title, diagram in [
    ('Recargas',
     'en_validacion → aprobada | rechazada'),
    ('Notificaciones',
     'pendiente → enviada → leida\n          → cancelada (auto-limpieza)\n          → fallida'),
    ('Solicitudes de Recarga',
     'pendiente → parcial → cumplida\n          → vencida\n          → cancelada'),
    ('Facturas (v3 con validacion_estado separado)',
     'pendiente (sin_validar) → pendiente (validada) → pagada\n                       → pendiente (rechazada) → anulada'),
    ('Obligaciones',
     'activa → en_progreso → completada\n       → cancelada'),
    ('Pagos',
     'en_proceso → pagado | fallido'),
]:
    h3(doc, title)
    code_block(doc, diagram, 'TEXT')

doc.add_page_break()

# ── PÁGINA FINAL ──────────────────────────────────────────────────────────────
add_blank(doc, 60)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Fin del Documento')
r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = NAVY_DARK

p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(6)
r2 = p2.add_run('━━━━━━━')
r2.font.color.rgb = GREEN; r2.font.bold = True; r2.font.size = Pt(12)

add_blank(doc, 14)

for line, italic in [
    ('DeOne Backend API v3.0', False),
    ('17 Endpoints Bot  |  2 Cron Jobs  |  4 Módulos', False),
    ('', False),
    ('7 de mayo de 2026', True),
    ('https://prueba-supabase.onrender.com/api', True),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(11); r.font.italic = italic
    r.font.color.rgb = GRAY if italic else NAVY_DARK

# ── Guardar ───────────────────────────────────────────────────────────────────
output = '/Users/geinermartinezmoscoso/Desktop/deone-oficial/deone-backend/docs/BOT_ENDPOINTS_v3.docx'
doc.save(output)
print(f'✅ Documento generado: {output}')
